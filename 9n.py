'''
    закидываем скриншоты в ворд файлы:
    в excel должны быть столбцы:
    "Номер скрина" == имя jpg файла (без типа файла и пути) "12345"
    "Новая цена" - если 0, то не учитывается (вроде бы)

    0 принимает стандартную рабочую таблицу
    1 чтение рабочей таблицы, отбираем строки с ценами и скринами
    2 обработка информации:
        объединяем скрины по компаниям и типам источников
    3 создание папок для хранения .doc файлов
    4 создание .doc файлов

    пример названия файла.doc:
    "1234_СИТИЛИНК ООО_03 компьютерное оборудование.doc"
'''
from excel_funcs import excel_to_list
import nine_editor

import os
import re
import docx
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Mm, Pt

# - - - * - -  НАСТРОЙКИ ЕКСЕЛЬ ФАЙЛА - - - * - - -
FOLDER_PATH = 'C:/Users/G.Tishchenko/Desktop/файлы_сэд/'
EXCEL_FOLDER = 'C:/Users/G.Tishchenko/Desktop/1 кв 2025/'
SCREEN_FOLDER_FOLDER = 'Z:/Тищенко Г.Л/2025_1 Скрины/'

# EXCEL_NAME = '26 Оборудование для театрально.xlsx'
# SCREEN_FOLDER_NAME = '26 театральное оборудование'
# EXCEL_NAME = '19 Бытовые приборы.xlsx'
# SCREEN_FOLDER_NAME = '19 бытовое оборудование'
# EXCEL_NAME = '3 компьютерное.xlsx'
# SCREEN_FOLDER_NAME = '3 компьютерное оборудование'
# EXCEL_NAME = '3 Нормирование.xlsx'
# SCREEN_FOLDER_NAME = 'нормирование'





part = '3 компьютерное'
# part = '26 Оборудование для театрально'
# part = '19 бытовые приборы'
# part = 'Нормирование'

# EXCEL_NAME = 'Проц 17' + '.xlsx'
EXCEL_NAME = part + '.xlsx'
SCREEN_FOLDER_NAME = part
EXCEL_PATH = EXCEL_FOLDER + EXCEL_NAME
SCREEN_FOLDER = SCREEN_FOLDER_FOLDER + SCREEN_FOLDER_NAME + "/"



ex_set = {
    'sheet_name': 'Лист1',
    'headers': False,
    'headers_names': [
        'Наименование ККН',
        'Источник ценовой информации',
        'ИНН поставщика',
        'Наименование поставщика',
        'Номер скрина',
        'Новая цена',
    ]
}

def edit_name(comp_name):

    re_sult = re.findall(r'\w+', comp_name.upper())
    return "_".join(re_sult)


def create_folders():
    try:
        os.mkdir(FOLDER_PATH)
    except:
        print('Главные папка уже есть')
    try:
        os.mkdir(FOLDER_PATH + 'Ответы на запрос')
    except:
        print('Ответы на запрос - папка уже есть')
    try:
        os.mkdir(FOLDER_PATH + 'Экранные копии')
    except:
        print('Экранные копии - папка уже есть')

def create_doc_test(file_path, comp_info):
    page_header = comp_info['source'].split(" ")[-1] + ' ' + comp_info['company_name']
    # page_header = page_header
    images = comp_info['screens']
    print("\n\nФайл:", file_path)
    print('колонтитул:', page_header,)
    for image in images:
        image_path = SCREEN_FOLDER + str(image['name']) + '.jpg'
        print("подпись", image['name'])
        print("Скрин", image_path)

def create_doc(file_path, comp_info):
    images = comp_info['screens']

    # Создаем word файл
    document = docx.Document()
    section = document.sections[0]
    # Горизонтальный формат листа
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    # отступы на странице
    section.top_margin = Inches(0.75)
    section.left_margin = Inches(0.7)
    # добавление верхнего колониткула
    header = section.header
    paragraph = header.paragraphs[0]
    page_header = comp_info['source'].split(" ")[-1] + ' ' + comp_info['company_name']
    paragraph.text = page_header

    # Закидываем в word файл скриншоты
    pic_number = 0
    paragraph_format = document.styles['Normal'].paragraph_format
    # paragraph_format.line_spacing = Pt(0) # междустрочный интервал
    paragraph_format.space_after  = Pt(0)
    try:
        for image in images:
            pic_number += 1
            # имя существующего скриншота
            image_path = SCREEN_FOLDER + str(image['name']) + '.jpg'
            image_exists = os.path.exists(image_path)
            if image_exists:
                document.add_paragraph(str(image['name']))
                document.add_picture(image_path, width = docx.shared.Inches(9.3), height = docx.shared.Inches(6.5))
                pic_number += 1
            else:
                print(image_path, "Не существует")

        document.save(file_path)
        print(f'{file_path} - сохранен, скриншотов: {pic_number}')
    except:
        print(f' ОШИБКААА : {page_header}')


def save_screens(information, logs = True):
    screen_amount_all = 0
    company_amount_all = 0

    for key, val in information.items():
        print('Определяем имя папки:', key, end = '\n\n') if logs else None
        screen_amount = 0
        company_amount = 0
        for company_inn, comp_info in val.items():
            file_name = FOLDER_PATH + key + '/' + comp_info['number'] + '_'+ edit_name(comp_info['company_name']) + '_' + edit_name(SCREEN_FOLDER_NAME) + '.docx'
            print('\n', file_name, comp_info) if logs else None
            company_amount += 1
            company_amount_all += 1
            screen_amount_cur = 0

            create_doc(file_name, comp_info)
            # create_doc_test(file_name, comp_info)


create_folders()
ex_list = excel_to_list(EXCEL_PATH, **ex_set)
# information = combine_info(ex_list)
information = nine_editor.combine_info_v2(ex_list)

save_screens(information, logs = False)
