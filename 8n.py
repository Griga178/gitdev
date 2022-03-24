# РАБОТА С DOC
'''
    Добавление фоток (.jpg) в word
    Создать папки: images, otveti, ekranki --- create_dirs()
        (куда будем сохранять word с Скринами)

    Из папки ales_screenes Otveti,Ekranki
                            ----- pic_save(dir_source, dir_destiny)

    1 Меняем
        screen_folder
'''
import docx
from docx.enum.section import WD_ORIENT
import os
from docx.shared import Inches

#name_file = 'new_file.docx'
#name_name_file = 'new_file_2.docx'

screen_folder = 'C:/Users/G.Tishchenko/Desktop/Screen_folder'

word_folder = 'C:/Users/G.Tishchenko/Desktop/'
dir_name_i = word_folder + 'Word_folder'
dir_name_O = dir_name_i + '/Otveti'
dir_name_E = dir_name_i + '/Ekranki'

def create_dirs():
    try:
        os.mkdir(dir_name_i)
        os.mkdir(dir_name_O)
        os.mkdir(dir_name_E)
    except:
        print('Папки не создавалась')




source_O = screen_folder + '/Otveti/'
source_E = screen_folder + '/Ekranki'

def pic_save(dir_source, dir_destiny):
    ''' 2 обязательных аргумента (названия папок)
    1) где лежат фотки
    2) куда сохраняем ворд файл
    '''
    all_pic_number = 0
    doc_number = 0
    list_comp_name = os.listdir(dir_source)
    for el in list_comp_name:
        # имя ворд файла
        doc_file_name = dir_destiny + '/' + el + '.docx'
        # Создаем word файл
        document = docx.Document()
        section = document.sections[0]
        # Горизонтальный формат листа
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        # отступы на странице
        section.top_margin = Inches(0.75)
        section.left_margin = Inches(0.7)
        # добавление верхнего колониткула
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = el

        #print(doc_file_name)

        # Закидываем в word файл скриншоты из папки
        pic_number = 0

        try:
            for image in (os.listdir(dir_source + '/' + el)):
                pic_number += 1
                # имя существующего скриншота
                name_image = dir_source + '/' + el + '/' +  image
                #print(page_number, name_image)
                document.add_picture(name_image, width=docx.shared.Inches(9.3), height=docx.shared.Inches(6.5))

            document.save(doc_file_name)
            all_pic_number += pic_number
            doc_number += 1
            print(f'{doc_file_name} - сохранен, скриншотов: {pic_number}')
        except:
            print(f' Нарушен стандарт в: {el}')

    print(f'\nВсего скриншотов: {all_pic_number}\nВсего word файлов: {doc_number}')

create_dirs()
pic_save(source_O, dir_name_O) # 691 шт 4 кв Ответы
pic_save(source_E, dir_name_E) # 745 шт 4 кв Экранки


def one_pic_save(dir_source, dir_destiny, comp_name):
    ''' 2 обязательных аргумента (названия папок)
    1) где лежат фотки
    2) куда сохраняем ворд файл
    3) имя компании
    '''
    all_pic_number = 0
    doc_number = 0

    list_comp_name = os.listdir(dir_source)
    for el in list_comp_name:
        if comp_name == el:
            # имя ворд файла
            doc_file_name = dir_destiny + '/' + el + '.docx'
            # Создаем word файл
            document = docx.Document()
            section = document.sections[0]
            # Горизонтальный формат листа
            new_width, new_height = section.page_height, section.page_width
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = new_width
            section.page_height = new_height
            # отступы на странице
            section.top_margin = Inches(0.75)
            section.left_margin = Inches(0.7)
            # добавление верхнего колониткула
            header = section.header
            paragraph = header.paragraphs[0]
            paragraph.text = el

            #print(doc_file_name)

            # Закидываем в word файл скриншоты из папки
            pic_number = 0

            #print(comp_name)
            try:
                for image in (os.listdir(dir_source + '/' + el)):
                    pic_number += 1
                    # имя существующего скриншота
                    name_image = dir_source + '/' + el + '/' +  image
                    #print(page_number, name_image)
                    document.add_picture(name_image, width=docx.shared.Inches(9.3), height=docx.shared.Inches(6.5))

                document.save(doc_file_name)
                all_pic_number += pic_number
                doc_number += 1
                print(f'{comp_name} - сохранен, скриншотов: {pic_number}')
            except:
                print(f' Нарушен стандарт в: {el}')

    print(f'\nВсего скриншотов: {all_pic_number}\nВсего word файлов: {doc_number}')


#one_file = 'СИТИЛИНК ООО'
#one_file = 'ОЛДИ ЛТД ЗАО'
#one_file = 'ОНЛАЙН ТРЕЙД ООО'
#one_file = 'ДНС РИТЕЙЛ ООО'
#one_file = 'М-ИНВЕСТ ООО'
#one_file = 'КОМПСЕРВИС ООО'
#one_file = 'НЕРА-М ООО'

#one_pic_save(source_O, dir_name_O, one_file)

#one_pic_save(source_E, dir_name_E, one_file)
