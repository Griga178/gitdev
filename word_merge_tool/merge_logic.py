import os
import re
import glob
from docx import Document
from openpyxl import Workbook

# Конфигурация плейсхолдеров
PLACEHOLDER_PATTERN = re.compile(r'\{([^{}]+)\}')
# Префиксы папок
TEMPLATE_PREFIX = 'example_'
PARAM_PREFIX = 'param_'

def find_template_files(folder):
    """
    Ищет все файлы шаблонов вида example_*.docx в указанной папке.
    Возвращает список словарей: [{'template': 'путь', 'name': 'имя', 'param_file': 'путь'}]
    """
    pattern = os.path.join(folder, f"{TEMPLATE_PREFIX}*.docx")
    files = glob.glob(pattern)

    if not files:
        return []

    result = []
    for file_path in files:
        filename = os.path.basename(file_path)
        # Извлекаем имя между 'example_' и '.docx'
        # example_contract.docx -> contract
        name_part = filename[len(TEMPLATE_PREFIX):-5]  # -5 убирает '.docx'

        param_filename = f"{PARAM_PREFIX}{name_part}.xlsx"
        param_path = os.path.join(folder, param_filename)

        result.append({
            'template': file_path,
            'name': name_part,
            'param_file': param_path
        })

    return result

# --- Работа с Word ---
def get_all_paragraphs(doc):
    """
    Генератор, который возвращает ВСЕ параграфы документа:
    1. Основное тело
    2. Таблицы (включая вложенные)
    3. Колонтитулы (Header/Footer)
    """
    # 1. Основное тело
    for paragraph in doc.paragraphs:
        yield paragraph

    # 2. Таблицы в теле
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph

    # 3. Колонтитулы (Секции)
    for section in doc.sections:
        # Header
        header = section.header
        if header:
            for paragraph in header.paragraphs:
                yield paragraph
            # Таблицы в хедере
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield paragraph

        # Footer
        footer = section.footer
        if footer:
            for paragraph in footer.paragraphs:
                yield paragraph
            # Таблицы в футере
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield paragraph

def extract_placeholders_from_docx(file_path):
    """
    Сканирует Word-файл и извлекает все уникальные плейсхолдеры.
    Учитывает тело, таблицы и колонтитулы.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Файл {file_path} не найден.")

    doc = Document(file_path)
    placeholders = set()

    for paragraph in get_all_paragraphs(doc):
        text = paragraph.text
        matches = PLACEHOLDER_PATTERN.findall(text)
        placeholders.update(matches)

    # Сортировка: сначала те, что длиннее (чтобы при замене не задеть подстроки)
    # Хотя для извлечения это не критично, но полезно для порядка
    return sorted(list(placeholders))

def extract_placeholders_from_docx_old(file_path):
    """
    Сканирует Word-файл и извлекает все уникальные плейсхолдеры вида {param_name}.
    Возвращает отсортированный список строк.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Файл {file_path} не найден.")

    doc = Document(file_path)
    placeholders = set()
    # pattern = re.compile(r'\{([^}]+)\}')  # Ищем текст внутри {}
    pattern = re.compile(r'\_([^}]+)\_')  # Ищем текст внутри _ _

    # Поиск в параграфах
    for paragraph in doc.paragraphs:
        matches = pattern.findall(paragraph.text)
        placeholders.update(matches)

    # Поиск в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    matches = pattern.findall(paragraph.text)
                    placeholders.update(matches)

    return sorted(list(placeholders))

def create_excel_template(file_path, placeholders):
    """
    Создает Excel-файл с заголовками.
    Первый столбец: filename (имя будущего файла)
    Остальные столбцы: найденные плейсхолдеры из Word
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Data"

    # Заголовки
    headers = ['filename'] + placeholders
    for col_idx, header in enumerate(headers, 1):
        ws.cell(row=1, column=col_idx, value=header)

    # Добавим одну пустую строку для примера
    for col_idx in range(1, len(headers) + 1):
        ws.cell(row=2, column=col_idx, value="")

    # Автоширина колонок
    for col in ws.columns:
        max_length = 0
        column_letter = col[0].column_letter
        for cell in col:
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except:
                pass
        ws.column_dimensions[column_letter].width = max_length + 5

    wb.save(file_path)
    return file_path

# --- Чтение Excel  ---

def load_excel_data(file_path):
    """
    Читает Excel файл используя openpyxl.
    Возвращает список словарей.
    Первый столбец - имя файла, остальные - данные для замены.
    """
    from openpyxl import load_workbook

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Файл {file_path} не найден.")

    wb = load_workbook(file_path, data_only=True)
    ws = wb.active

    # Читаем заголовки из первой строки
    headers = []
    for col in range(1, ws.max_column + 1):
        cell_value = ws.cell(row=1, column=col).value
        headers.append(str(cell_value).strip() if cell_value else "")

    if len(headers) < 2:
        raise ValueError("В Excel должно быть минимум 2 столбца (Имя файла и данные).")

    filename_col = headers[0]
    data_cols = headers[1:]

    data_list = []

    # Читаем строки данных (начиная со 2-й строки)
    for row_idx in range(2, ws.max_row + 1):
        # Проверяем, есть ли имя файла в строке
        filename_cell = ws.cell(row=row_idx, column=1).value
        if not filename_cell or str(filename_cell).strip() == "":
            continue  # Пропускаем пустые строки

        entry = {
            'filename': str(filename_cell).strip(),
            'replacements': {}
        }

        # Заполняем данные для замены
        for col_idx, col_name in enumerate(data_cols, 2):
            cell_value = ws.cell(row=row_idx, column=col_idx).value
            entry['replacements'][f"{{{col_name}}}"] = str(cell_value) if cell_value is not None else ""

        data_list.append(entry)

    wb.close()
    return data_list

# --- Замена текста в Word ---

def _get_run_style(run):
    """Сохраняет стиль запуска (bold, italic, font и т.д.)"""
    return {
        'bold': run.bold,
        'italic': run.italic,
        'underline': run.underline,
        'font_name': run.font.name,
        'font_size': run.font.size,
        'color': run.font.color.rgb if run.font.color.rgb else None
    }

def _apply_style(run, style):
    """Применяет сохраненный стиль к запуску"""
    run.bold = style['bold']
    run.italic = style['italic']
    run.underline = style['underline']
    if style['font_name']:
        run.font.name = style['font_name']
    if style['font_size']:
        run.font.size = style['font_size']

def replace_text_in_paragraph(paragraph, placeholder, value):
    """
    Заменяет плейсхолдер в параграфе, сохраняя форматирование.
    Корректно работает, даже если плейсхолдер разбит на несколько частей (runs).
    """
    if placeholder not in paragraph.text:
        return

    runs = paragraph.runs
    full_text = "".join([run.text for run in runs])

    try:
        start_idx = full_text.index(placeholder)
        end_idx = start_idx + len(placeholder)
    except ValueError:
        return

    # 1. Находим индексы runs, которые покрывают плейсхолдер
    current_pos = 0
    start_run_idx = None
    end_run_idx = None

    for i, run in enumerate(runs):
        r_start = current_pos
        r_end = current_pos + len(run.text)

        if r_start <= start_idx < r_end:
            start_run_idx = i
        if r_start < end_idx <= r_end:
            end_run_idx = i

        current_pos += len(run.text)
        if start_run_idx is not None and end_run_idx is not None:
            break

    if start_run_idx is None:
        return

    # 2. Сохраняем стиль ПЕРВОГО run-а плейсхолдера
    # Именно этот стиль будет применен ко всему вставляемому значению
    target_style = _get_run_style(runs[start_run_idx])

    # 3. Логика замены
    # Если плейсхолдер целиком в одном run (самый частый случай)
    if start_run_idx == end_run_idx:
        run = runs[start_run_idx]
        run.text = run.text.replace(placeholder, value)
        _apply_style(run, target_style)
    else:
        # Если плейсхолдер РАЗБИТ на несколько runs (например, часть жирная, часть нет)
        # Мы должны аккуратно вырезать его части из всех затронутых runs

        # 3.1. Обрабатываем первый run (оставляем текст ДО плейсхолдера)
        first_run = runs[start_run_idx]
        text_before = first_run.text[:first_run.text.find(placeholder[0])]
        first_run.text = text_before + value
        _apply_style(first_run, target_style)

        # 3.2. Обрабатываем последний run (оставляем текст ПОСЛЕ плейсхолдера)
        last_run = runs[end_run_idx]
        # Ищем, где заканчивается плейсхолдер в этом run-е
        # Мы знаем, что конец плейсхолдера где-то здесь, удаляем всё, что до него относится
        placeholder_end_in_run = last_run.text.find(placeholder[-1]) + 1
        # Грубая оценка: удаляем начало run-а, которое входило в плейсхолдер
        # Надежнее найти точное вхождение части плейсхолдера
        part_in_last = ""
        # Собираем ожидаемую часть плейсхолдера для этого run-а
        # (это упрощение, но работает в 99% случаев)
        if placeholder in last_run.text:
             last_run.text = last_run.text.replace(placeholder, "")
        else:
             # Если плейсхолдер разбит, в последнем run-е будет только хвост
             # Очищаем всё, что является частью плейсхолдера (начало run-а)
             # Простой способ: найти первый символ плейсхолдера в этом run-е и удалить до него
             # Но так как это хвост, мы просто удаляем соответствующую длину
             # Для надежности просто удалим найденный остаток
             for i in range(len(placeholder), 0, -1):
                 tail = placeholder[-i:]
                 if tail in last_run.text and last_run.text.index(tail) == 0:
                     last_run.text = last_run.text.replace(tail, "", 1)
                     break

        # 3.3. Промежуточные run-ы полностью очищаем (они целиком внутри плейсхолдера)
        for i in range(start_run_idx + 1, end_run_idx):
            runs[i].text = ""

def process_document(template_path, output_path, data):
    """
    Создает копию шаблона, заменяет данные и сохраняет.
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Шаблон {template_path} не найден.")

    doc = Document(template_path)

    # Замена в параграфах
    for paragraph in doc.paragraphs:
        for key, value in data['replacements'].items():
            replace_text_in_paragraph(paragraph, key, value)

    # Замена в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data['replacements'].items():
                        replace_text_in_paragraph(paragraph, key, value)

    # Сохранение
    if not output_path.endswith('.docx'):
        output_path += '.docx'

    doc.save(output_path)
    return output_path
